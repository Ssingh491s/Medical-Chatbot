from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import shutil

# Create a new Word document
doc = Document()

# Title Style
title = doc.add_heading("AI-Powered Chatbot using DeepSeek API with Tkinter GUI", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Section: Project Overview
doc.add_heading("Project Overview", level=1)
doc.add_paragraph(
    "This project implements a desktop-based chatbot application using the Tkinter GUI "
    "library in Python, integrated with the DeepSeek large language model (LLM) API. "
    "The bot accepts user queries through a friendly interface and responds with "
    "intelligent, context-aware replies."
)

# Section: Technologies Used
doc.add_heading("Technologies Used", level=1)
doc.add_paragraph("• Language: Python")
doc.add_paragraph("• GUI: Tkinter")
doc.add_paragraph("• API: DeepSeek Chat Completions API")
doc.add_paragraph("• Libraries: requests, json, tkinter, ttk, scrolledtext, threading")

# Section: Key Features
doc.add_heading("Key Features", level=1)
doc.add_paragraph("• Interactive and scrollable chat window")
doc.add_paragraph("• Real-time messaging with Enter/Send button")
doc.add_paragraph("• Asynchronous API call handling via threads (prevents UI freezing)")
doc.add_paragraph("• Stylized chat output for user, bot, and error responses")
doc.add_paragraph("• \"Clear Chat\" functionality to reset the conversation")

# Section: Sample GUI Screenshot
doc.add_heading("Sample GUI Screenshot", level=1)
screenshot_path = "/mnt/data/chatbot_gui_sample.jpg"
# Replace this path with the actual screenshot if available
shutil.copy("/mnt/data/image_0.jpg", screenshot_path)  # Use previously uploaded image
doc.add_picture(screenshot_path, width=Inches(4.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Section: Program Flow Diagram
doc.add_heading("Program Flow Diagram", level=1)
diagram_text = (
    "+----------------------+\n"
    "|      Start GUI       |\n"
    "+----------+-----------+\n"
    "           |\n"
    "           v\n"
    "+----------------------+\n"
    "|  User Enters Prompt  |\n"
    "+----------+-----------+\n"
    "           |\n"
    "           v\n"
    "+-------------------------------+\n"
    "|  Display Message in Chat Box |\n"
    "+----------+--------------------+\n"
    "           |\n"
    "           v\n"
    "+-------------------------------+\n"
    "|  Send Prompt to DeepSeek API |\n"
    "+-------------------------------+\n"
    "           |\n"
    "           v\n"
    "+-------------------------------+\n"
    "|  Receive & Display Response  |\n"
    "+-------------------------------+\n"
    "           |\n"
    "           v\n"
    "+----------------------+\n"
    "|   Wait for New Input |\n"
    "+----------------------+"
)
doc.add_paragraph(diagram_text, style='Intense Quote')

# Section: Security Consideration
doc.add_heading("Security Consideration", level=1)
doc.add_paragraph(
    "⚠️ API Key is hardcoded. In production, store it securely using environment variables "
    "or a secrets manager to avoid leaking credentials."
)

# Section: Sample Output
doc.add_heading("Sample Output", level=1)
doc.add_paragraph("You: Hello, how are you?\nchatbot: I'm just a virtual assistant, but I'm here to help you. How can I assist you today?")

# Save the document
doc_path = r"C:\Users\shubh\Downloads\chatbot-python-project-data-codes (1)\python Api deep\AI_Chatbot_Report.docx"
doc.save(doc_path)

doc_path



